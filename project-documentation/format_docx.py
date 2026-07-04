
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

input_docx = "project-documentation/Chapter-1-SPMP.docx"
output_docx = "project-documentation/Chapter-1-SPMP-formatted.docx"
doc = Document(input_docx)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Apply heading styles and spacing
for para in doc.paragraphs:
    text = para.text.strip()
    if text.startswith('1.') or text.startswith('2.') or text.startswith('3.') or text.startswith('4.') or text.startswith('5.') or text.startswith('6.'):
        para.style = doc.styles['Heading 1']
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_after = Pt(18)
    elif text.startswith('1.1') or text.startswith('2.1') or text.startswith('3.1') or text.startswith('4.1') or text.startswith('5.1'):
        para.style = doc.styles['Heading 2']
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.space_after = Pt(12)
    elif text.startswith('1.1.1') or text.startswith('2.1.1'):
        para.style = doc.styles['Heading 3']
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.space_after = Pt(8)
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.space_after = Pt(6)

# Set table style (fallback if 'Table Grid' is missing)
def set_table_style(table):
    try:
        table.style = 'Table Grid'
    except KeyError:
        try:
            table.style = 'Normal Table'
        except KeyError:
            pass  # Use whatever style is available
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for table in doc.tables:
    set_table_style(table)

# Add page breaks before major sections
for para in doc.paragraphs:
    if para.style.name == 'Heading 1' and para.text.strip() != doc.paragraphs[0].text.strip():
        run = para.insert_paragraph_before().add_run()
        run.add_break()

# Save the improved document
doc.save(output_docx)
print(f"Formatted and saved as {output_docx}")
